from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from datetime import datetime
import shutil
import sqlite3
import os
import uuid
import io
import httpx

app = FastAPI(title="PROCON Galvão - Gerador de CP")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# URL do Gotenberg — configurada via variável de ambiente no Render
GOTENBERG_URL = os.getenv("GOTENBERG_URL", "http://localhost:3000")

DB_PATH = "documentos.db"
UPLOAD_DIR = "uploads"

os.makedirs(UPLOAD_DIR, exist_ok=True)


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_db() as conn:
        conn.execute('''CREATE TABLE IF NOT EXISTS modelos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL UNIQUE,
            arquivo TEXT NOT NULL
        )''')
        conn.commit()


init_db()


def formatar_data_por_extenso():
    meses = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
    }
    hoje = datetime.today()
    return f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"


def substituir_texto(doc, placeholder, novo_texto):
    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, novo_texto)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, novo_texto)


@app.get("/modelos")
def listar_modelos():
    with get_db() as conn:
        rows = conn.execute("SELECT id, nome FROM modelos").fetchall()
    return [{"id": r["id"], "nome": r["nome"]} for r in rows]


@app.post("/modelos")
async def adicionar_modelo(arquivo: UploadFile = File(...)):
    if not arquivo.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Apenas arquivos .docx são aceitos.")

    nome_modelo = arquivo.filename.replace(".docx", "")
    unique_name = f"{uuid.uuid4()}_{arquivo.filename}"
    caminho = os.path.join(UPLOAD_DIR, unique_name)

    with open(caminho, "wb") as f:
        shutil.copyfileobj(arquivo.file, f)

    try:
        with get_db() as conn:
            conn.execute(
                "INSERT OR REPLACE INTO modelos (nome, arquivo) VALUES (?, ?)",
                (nome_modelo, caminho)
            )
            conn.commit()
    except Exception as e:
        os.remove(caminho)
        raise HTTPException(status_code=500, detail=str(e))

    return {"mensagem": "Modelo adicionado com sucesso!", "nome": nome_modelo}


@app.delete("/modelos/{nome}")
def remover_modelo(nome: str):
    with get_db() as conn:
        row = conn.execute("SELECT arquivo FROM modelos WHERE nome = ?", (nome,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Modelo não encontrado.")
        try:
            if os.path.exists(row["arquivo"]):
                os.remove(row["arquivo"])
        except Exception:
            pass
        conn.execute("DELETE FROM modelos WHERE nome = ?", (nome,))
        conn.commit()
    return {"mensagem": "Modelo removido com sucesso!"}


@app.post("/gerar")
async def gerar_documento(
    modelo: str = Form(...),
    polo_ativo: str = Form(...),
    numero_reclamacao: str = Form(...),
    comarca: str = Form("")
):
    with get_db() as conn:
        row = conn.execute("SELECT arquivo FROM modelos WHERE nome = ?", (modelo,)).fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Modelo não encontrado.")

    caminho_arquivo = row["arquivo"]
    if not os.path.exists(caminho_arquivo):
        raise HTTPException(status_code=500, detail="Arquivo do modelo não encontrado no servidor.")

    doc = Document(caminho_arquivo)
    substituicoes = {
        "POLOATIVO": polo_ativo,
        "NRECLAMAÇÃO": numero_reclamacao,
        "DATAHOJE": formatar_data_por_extenso(),
        "CMJUIZO": comarca,
    }
    for chave, valor in substituicoes.items():
        substituir_texto(doc, chave, valor)

    # Salva o .docx preenchido em memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Envia para o Gotenberg converter para PDF
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(
                f"{GOTENBERG_URL}/forms/libreoffice/convert",
                files={"files": ("documento.docx", buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail=f"Erro na conversão: {response.text}")
    except httpx.ConnectError:
        raise HTTPException(status_code=500, detail="Serviço de conversão indisponível. Tente novamente em instantes.")
    except httpx.TimeoutException:
        raise HTTPException(status_code=500, detail="Tempo esgotado na conversão. Tente novamente.")

    nome_final = f"{polo_ativo} - CARTA DE PREPOSIÇÃO.pdf"
    return StreamingResponse(
        io.BytesIO(response.content),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{nome_final}"'}
    )


@app.get("/health")
def health():
    return {"status": "ok"}
